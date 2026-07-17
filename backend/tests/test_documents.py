import io
import docx
from fastapi.testclient import TestClient
from app.main import app
from app.config.settings import settings

client = TestClient(app)

def create_mock_docx() -> bytes:
    """Helper to generate a valid minimal DOCX in memory."""
    doc = docx.Document()
    doc.add_paragraph("This is a paragraph from a mock Word document.")
    doc.add_paragraph("Another paragraph with test data.")
    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()

def test_upload_txt_file():
    # 1. Upload Plain Text
    content = b"This is clean text from a patient note."
    files = {"file": ("test_note.txt", content, "text/plain")}
    response = client.post("/documents/upload", files=files)
    
    assert response.status_code == 201
    json_data = response.json()
    assert "document" in json_data
    assert json_data["document"]["filename"] == "test_note.txt"
    assert json_data["document"]["type"] == "txt"
    assert json_data["document"]["status"] == "completed"
    assert json_data["extracted_text"] == "This is clean text from a patient note."
    assert "debug" in json_data
    assert json_data["debug"]["execution_time_ms"] > 0
    assert len(json_data["debug"]["logs"]) > 0

    doc_id = json_data["document"]["id"]

    # 2. Get Single Document Details
    get_response = client.get(f"/documents/{doc_id}")
    assert get_response.status_code == 200
    get_data = get_response.json()
    assert get_data["document"]["id"] == doc_id
    assert get_data["extracted_text"] == "This is clean text from a patient note."

    # 3. List All Documents
    list_response = client.get("/documents")
    assert list_response.status_code == 200
    list_data = list_response.json()
    assert len(list_data["documents"]) >= 1
    assert any(d["id"] == doc_id for d in list_data["documents"])

    # 4. Delete Document
    del_response = client.delete(f"/documents/{doc_id}")
    assert del_response.status_code == 200
    assert del_response.json()["status"] == "deleted"

    # 5. Confirm deletion
    get_fail = client.get(f"/documents/{doc_id}")
    assert get_fail.status_code == 404

def test_upload_docx_file():
    docx_bytes = create_mock_docx()
    files = {"file": ("test_file.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    response = client.post("/documents/upload", files=files)
    
    assert response.status_code == 201
    json_data = response.json()
    assert json_data["document"]["filename"] == "test_file.docx"
    assert json_data["document"]["type"] == "docx"
    assert "mock Word document" in json_data["extracted_text"]

    doc_id = json_data["document"]["id"]
    
    # Cleanup
    client.delete(f"/documents/{doc_id}")

def test_upload_unsupported_file():
    files = {"file": ("script.py", b"print('hello')", "text/x-python")}
    response = client.post("/documents/upload", files=files)
    assert response.status_code == 400
    assert "Unsupported file type" in response.json()["detail"]["error"]
