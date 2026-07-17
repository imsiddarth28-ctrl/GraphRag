import os
import time
import uuid
from datetime import datetime, timezone
from typing import Tuple
import pdfplumber
import docx
from app.models.document import Document
from app.repositories.document_repository import DocumentRepository

class UnsupportedFileTypeError(Exception):
    """Raised when an uploaded file type is not supported."""
    pass

class DocumentService:
    """
    Manages document uploads, raw text extraction, and metadata persistence.
    """
    def __init__(self, repository: DocumentRepository, upload_dir: str = "data/uploads"):
        self.repository = repository
        self.upload_dir = upload_dir
        os.makedirs(self.upload_dir, exist_ok=True)

    async def upload_document(self, filename: str, content: bytes, size: int) -> Tuple[Document, str, float]:
        """
        Coordinates document persistence, text extraction, database storage, and timing.
        Returns (Document, extracted_text, processing_time_ms).
        """
        start_time = time.perf_counter()
        
        # Determine file type
        ext = os.path.splitext(filename)[1].lower()
        if ext not in [".txt", ".pdf", ".docx", ".md"]:
            raise UnsupportedFileTypeError(f"Unsupported file type '{ext}'. Supported: .txt, .pdf, .docx, .md")

        # Generate unique ID
        doc_id = str(uuid.uuid4())
        
        # Save file to disk
        file_path = os.path.join(self.upload_dir, f"{doc_id}{ext}")
        with open(file_path, "wb") as f:
            f.write(content)

        extracted_text = ""
        status = "completed"

        try:
            # Parse text contents
            if ext == ".txt" or ext == ".md":
                extracted_text = self._parse_text(file_path)
            elif ext == ".pdf":
                extracted_text = self._parse_pdf(file_path)
            elif ext == ".docx":
                extracted_text = self._parse_docx(file_path)
        except Exception as e:
            status = "failed"
            extracted_text = f"Parsing Error: {str(e)}"
            # Delete corrupted/failed file if writing failed
            if os.path.exists(file_path):
                os.remove(file_path)
            raise e

        # Create document domain model
        doc = Document(
            id=doc_id,
            filename=filename,
            type=ext.replace(".", ""),
            size=size,
            created_at=datetime.now(timezone.utc),
            status=status
        )

        # Save metadata and extracted text to SQLite Repository
        await self.repository.save(doc, extracted_text)

        processing_time_ms = (time.perf_counter() - start_time) * 1000.0
        return doc, extracted_text, processing_time_ms

    def _parse_text(self, path: str) -> str:
        """
        Reads plain text / markdown files.
        """
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    def _parse_pdf(self, path: str) -> str:
        """
        Extracts raw text pages from a PDF.
        """
        text_parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n".join(text_parts)

    def _parse_docx(self, path: str) -> str:
        """
        Extracts raw text paragraphs and tables from a Word document.
        """
        doc = docx.Document(path)
        text_parts = []
        
        # Process paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
                
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
                    
        return "\n".join(text_parts)
